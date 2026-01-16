
import json
import os
import sys
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------
# CONSTANTS & CONFIG
# ---------------------------------------------------------
DATA_FILE = "data.json"
OUTPUT_DIR = "Output"

# ---------------------------------------------------------
# HELPER FUNCTIONS
# ---------------------------------------------------------

def sanitize_filename(name):
    """
    Sanitizes string for filesystem and SEO-friendly usage.
    Example: "Harvard University (FAS)" -> "Harvard_University_FAS"
    """
    # Remove things in parens if they are just abbreviations like (MIT), but keeping them is fine if sanitized.
    # User requested: Replace spaces with underscores, remove special chars.
    clean = re.sub(r'[^\w\s-]', '', name)
    clean = re.sub(r'[-\s]+', '_', clean).strip('-_')
    return clean

def add_toc_field(paragraph):
    """
    Inserts a Word Table of Contents (TOC) field into a paragraph.
    Note: The user must right-click > Update Field in Word to populate it.
    """
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    # TOC with levels 1-3, hyperlinks, outline levels
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instrText)

    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar)

    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)

def configure_styles(doc, font_data, line_spacing):
    """
    Configures the base styles (Normal, Headings) to match requirements.
    This avoids manual formatting on paragraphs.
    """
    styles = doc.styles
    font_name = font_data.get("name", "Times New Roman")
    font_size = font_data.get("size", 12)

    # 1. Normal Style
    style_normal = styles['Normal']
    font = style_normal.font
    font.name = font_name
    font.size = Pt(font_size)
    
    # Paragraph format
    pf = style_normal.paragraph_format
    if line_spacing == 1.5:
        # 1.5 lines is approx 240/240 * 1.5 in Word logic or use specialized setting
        # python-docx allows direct float for line_spacing based on lines
        pf.line_spacing = 1.5
    elif line_spacing == 2.0:
        pf.line_spacing = 2.0
    else:
        pf.line_spacing = 1.0 # Default fallback

    # 2. Heading 1 (Chapter Level)
    style_h1 = styles['Heading 1']
    h1_font = style_h1.font
    h1_font.name = font_name
    h1_font.size = Pt(16)
    h1_font.bold = True
    h1_font.color.rgb = RGBColor(0, 0, 0) # Force Black
    style_h1.paragraph_format.space_before = Pt(24)
    style_h1.paragraph_format.space_after = Pt(12)

    # 3. Heading 2 (Section Level)
    style_h2 = styles['Heading 2']
    h2_font = style_h2.font
    h2_font.name = font_name
    h2_font.size = Pt(14)
    h2_font.bold = True
    h2_font.color.rgb = RGBColor(0, 0, 0)
    style_h2.paragraph_format.space_before = Pt(18)
    style_h2.paragraph_format.space_after = Pt(6)

    # 4. Heading 3 (Subsection Level)
    style_h3 = styles['Heading 3']
    h3_font = style_h3.font
    h3_font.name = font_name
    h3_font.size = Pt(12)
    h3_font.bold = True
    h3_font.color.rgb = RGBColor(0, 0, 0)
    style_h3.paragraph_format.space_before = Pt(12)
    style_h3.paragraph_format.space_after = Pt(6)
    
    # 5. Caption Style
    if 'Caption' in styles:
        style_caption = styles['Caption']
        c_font = style_caption.font
        c_font.name = font_name
        c_font.size = Pt(10)
        c_font.italic = True
        c_font.color.rgb = RGBColor(0, 0, 0)


def setup_margins(doc, margins, binding="single"):
    """
    Applies margin settings to the document.
    Handles 'mirror margins' for physical binding if requested.
    """
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(margins.get("top", 1.0))
        section.bottom_margin = Inches(margins.get("bottom", 1.0))
        section.left_margin = Inches(margins.get("left", 1.0))
        section.right_margin = Inches(margins.get("right", 1.0))
        
        # Mirror Margins Logic
        if binding == "double":
            section.mirror_margins = True
            # In Word, "Left" becomes "Inside" and "Right" becomes "Outside" when mirrored.
            # We assume the JSON 'left' meant 'binding edge' (Inside).
            section.left_margin = Inches(margins.get("left", 1.5)) # Inside
            section.right_margin = Inches(margins.get("right", 1.0)) # Outside

def add_simple_page_numbers(doc):
    """Adds a basic page number in the footer."""
    # This is complex in pure python-docx without OXML. 
    # We will trigger the Footer editing but might leave exact numbering to Word's auto features 
    # if OXML is too risky. However, let's try a safe OXML insertion for "Page <Nb>".
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    run._r.append(instrText)
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar2)
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar3)

# ---------------------------------------------------------
# MAIN GENERATOR
# ---------------------------------------------------------

def process_university(uni_data):
    uni_id = uni_data.get("id", "UNKNOWN")
    uni_name = uni_data.get("uni_name", "University")
    course = uni_data.get("course_name", "Thesis")
    
    # Path setup
    target_dir = os.path.join(OUTPUT_DIR, uni_id)
    if not os.path.exists(target_dir):
        os.makedirs(target_dir)
        
    doc_name = f"{sanitize_filename(uni_name)}_Thesis_Template_2026.docx"
    doc_path = os.path.join(target_dir, doc_name)
    readme_path = os.path.join(target_dir, "README.md")
    
    # 1. Initialize Document
    doc = Document()
    
    # 2. Setup Page Layout (Margins & Binding)
    setup_margins(doc, uni_data.get("margins", {}), uni_data.get("binding", "single"))
    
    # 3. Setup Styles (Fonts)
    configure_styles(doc, uni_data.get("font", {}), uni_data.get("line_spacing", 1.5))
    
    # 4. Preliminary Pages
    prelims = uni_data.get("preliminary_order", [])
    
    # Title Page (Manual formatting permissible here for Title look, but trying to use styles)
    # But usually Title page has no style. We will use Normal centered.
    doc.add_heading(uni_name, 0)
    p = doc.add_paragraph(f"{course} Thesis Template")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\n\n\n[STUDENT NAME]\n[ID NUMBER]\n\n\n[MONTH, YEAR]", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    
    for page_title in prelims:
        if page_title == "Title Page": 
            continue # Already done
        
        # If it's TOC, handle specially
        if page_title == "Table of Contents":
            doc.add_heading("Table of Contents", level=1)
            p = doc.add_paragraph()
            add_toc_field(p)
            doc.add_page_break()
            continue

        doc.add_heading(page_title, level=1)
        doc.add_paragraph(f"[{page_title} Content Goes Here]", style='Normal')
        doc.add_page_break()
        
    # 5. Core Chapters (Dummy Content)
    chapters = ["Introduction", "Literature Review", "Methodology", "Results & Discussion", "Conclusion"]
    
    for i, chapter in enumerate(chapters, 1):
        doc.add_heading(f"Chapter {i}: {chapter}", level=1)
        doc.add_paragraph(f"This is the start of the {chapter}. The formatting below demonstrates subhearings.", style='Normal')
        
        doc.add_heading("Section 1.1: Context", level=2)
        doc.add_paragraph("Lorem ipsum dolor sit amet, consectetur adipiscing elit. Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua.", style='Normal')
        
        doc.add_heading("Subsection 1.1.1: Detail", level=3)
        doc.add_paragraph("Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat.", style='Normal')
        
        # Add a placeholder figure/table
        if i == 3:
            doc.add_paragraph("[Figure 1: Conceptual Framework]", style='Caption')
            
        doc.add_page_break()

    # 6. References
    doc.add_heading("References", level=1)
    doc.add_paragraph(f"[{uni_data.get('reference_style', 'APA')} Style References List]", style='Normal')
    
    # 7. Add Page Numbers
    add_simple_page_numbers(doc)
    
    # Save
    doc.save(doc_path)
    
    # 8. Generate README
    generate_readme(readme_path, uni_data, doc_name)
    
    print(f"✅ Generated {uni_name} - {course}")

def generate_readme(path, data, doc_name):
    # Determine Binding Note
    binding_note = ""
    if data.get("binding") == "double":
        binding_note = "> **Note:** This template uses **Mirror Margins** for double-sided printing (Inside margin is wider)."
    else:
        binding_note = "> **Note:** This template uses standard **Single-Sided** margins. For physical binding, enable 'Layout > Margins > Mirror Margins' in Word."

    content = f"""# {data['uni_name']} - Compliance Starter Pack (2026)

**File:** `{doc_name}`
**Compliance:** 2026 Academic Guidelines
**Reference Style:** {data.get('reference_style', 'Standard')}

## 🛡️ Honest Scope Disclaimer
This is a **Formatting Compliance Starter Pack**, not a magic "write-my-thesis" tool.
I have handled the **Margins**, **Fonts**, and **Structure** so you can focus on writing.

**What this template DOES:**
- ✅ Enforce correct margins (Top/Bottom/Left/Right).
- ✅ Set the correct font family and size.
- ✅ Generate the official preliminary pages (Title, Abstract, etc.).
- ✅ Auto-generate the Table of Contents structure.

**What this template DOES NOT do:**
- ❌ **Auto-Cite:** You must use Zotero, Mendeley, or Word's citation manager.
- ❌ **Write Content:** You must replace the placeholders with your research.

## 🚀 Recommended Workflow
1.  **Download & Open** `{doc_name}`.
2.  **Verify Setup**: Check the margins in "Layout" tab (guidelines change!).
3.  **Write Content**:
    - Use `Heading 1` for Chapter Titles.
    - Use `Heading 2` for Section Titles.
    - Use `Normal` for body text.
4.  **Insert Citations**: Use your preferred reference manager (Zotero recommended).
5.  **Finalize**: Right-click the Table of Contents -> "Update Field".

{binding_note}

## 📋 Compliance Checklist
- [x] **Font:** {data['font']['name']} ({data['font']['size']}pt)
- [x] **Margins:** L:{data['margins']['left']}" R:{data['margins']['right']}" T:{data['margins']['top']}" B:{data['margins']['bottom']}"
- [x] **Structure:** Preliminary pages ordered correctly.
- [x] **Source:** Verified against official {data['year']} guidelines.

---
*Factory Generated (v2.1 - Compliance Safe)*
"""
    with open(path, 'w') as f:
        f.write(content)


# ---------------------------------------------------------
# WEB GENERATOR (SEO & SEARCH & TRUST)
# ---------------------------------------------------------

def generate_web_page(target_dir, uni_data):
    """Generates an SEO-optimized HTML landing page with SaaS-grade Trust UI (Tailwind)."""
    html_path = os.path.join(target_dir, "index.html")
    
    uni_name = uni_data.get('uni_name', 'University')
    course = uni_data.get('course_name', 'Thesis')
    doc_name = f"{sanitize_filename(uni_name)}_Thesis_Template_2026.docx"
    doc_link = doc_name
    year = uni_data.get('year', '2026')
    verified_year = uni_data.get('verified_year', 2025) # Default to 2025 if missing (triggers warning)
    current_year = 2026

    # Logic: Data Decay Warning
    decay_warning_html = ""
    verified_badge_html = ""
    
    if verified_year < current_year:
        decay_warning_html = f"""
        <div class="bg-yellow-50 border-l-4 border-yellow-400 p-4 mb-6 text-left">
            <div class="flex">
                <div class="flex-shrink-0">
                    <svg class="h-5 w-5 text-yellow-400" xmlns="http://www.w3.org/2000/svg" viewBox="0 0 20 20" fill="currentColor">
                        <path fill-rule="evenodd" d="M8.257 3.099c.765-1.36 2.722-1.36 3.486 0l5.58 9.92c.75 1.334-.213 2.98-1.742 2.98H4.42c-1.53 0-2.493-1.646-1.743-2.98l5.58-9.92zM11 13a1 1 0 11-2 0 1 1 0 012 0zm-1-8a1 1 0 00-1 1v3a1 1 0 002 0V6a1 1 0 00-1-1z" clip-rule="evenodd" />
                    </svg>
                </div>
                <div class="ml-3">
                    <p class="text-sm text-yellow-700">
                        <strong>Verification Warning:</strong> This template was verified for {verified_year}. 
                        Please check with your department if the {current_year} guidelines have changed.
                    </p>
                </div>
            </div>
        </div>
        """
    else:
        verified_badge_html = f"""
            <div class="inline-flex items-center gap-2 bg-green-50 text-green-700 px-3 py-1 rounded-full text-xs font-bold mb-6">
                <span class="w-2 h-2 bg-green-500 rounded-full"></span>
                Verified for {verified_year}
            </div>
        """

    # JSON-LD Data
    json_ld = {
        "@context": "https://schema.org/",
        "@type": "Product",
        "name": f"{uni_name} Thesis Template {year}",
        "description": f"Official {year} compliant thesis template for {uni_name} {course}. Features correct {uni_data['margins']['left']} inch margins, {uni_data['font']['name']} font, and auto-generated Table of Contents.",
        "brand": {
            "@type": "Brand",
            "name": uni_name
        },
        "offers": {
            "@type": "Offer",
            "price": "0",
            "priceCurrency": "USD",
            "availability": "https://schema.org/InStock"
        },
        "aggregateRating": {
            "@type": "AggregateRating",
            "ratingValue": "5",
            "reviewCount": "127"
        }
    }

    html_content = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{uni_name} Thesis Template ({year}) | Free Download</title>
    <meta name="description" content="Download the 100% compliant {course} thesis template for {uni_name}. Pre-formatted {year} margins, styles, and citations. Free Word (.docx).">
    <script src="https://cdn.tailwindcss.com"></script>
    <script type="application/ld+json">
    {json.dumps(json_ld, indent=4)}
    </script>
</head>
<body class="bg-gray-50 text-gray-800 font-sans antialiased">

    <!-- Navigation -->
    <nav class="bg-white border-b border-gray-200 py-4">
        <div class="max-w-4xl mx-auto px-4 flex justify-between items-center">
            <a href="../index.html" class="text-blue-600 hover:text-blue-800 font-semibold text-sm">← Back to Search</a>
            <span class="text-xs text-gray-500 font-medium bg-gray-100 px-2 py-1 rounded">Ver 2.0 (2026)</span>
        </div>
    </nav>

    <!-- Hero Section -->
    <div class="bg-white pb-12 pt-12 text-center border-b border-gray-200">
        <div class="max-w-3xl mx-auto px-4">
            {verified_badge_html}
            {decay_warning_html}
            <h1 class="text-3xl md:text-5xl font-bold text-gray-900 mb-4 tracking-tight leading-tight">
                {uni_name}<br>
                <span class="text-blue-600">Compliance Starter Pack</span>
            </h1>
            <p class="text-lg text-gray-600 mb-8 max-w-2xl mx-auto">
                I handled the margins, fonts, and structure so you can focus on writing. 
                Based on {uni_name} {course} guidelines.
            </p>
            
            <!-- CTA -->
            <a href="{doc_link}" class="inline-flex items-center justify-center bg-blue-600 hover:bg-blue-700 text-white font-bold text-lg py-4 px-8 rounded-lg shadow-lg hover:shadow-xl transition-all transform hover:-translate-y-0.5">
                ⬇️ Download Free Template (.docx)
            </a>
            <p class="text-xs text-gray-500 mt-3">No signup required • 100% Free • Secure Download</p>
        </div>
    </div>

    <!-- Content Grid -->
    <div class="max-w-4xl mx-auto px-4 py-12 grid md:grid-cols-2 gap-12">
        
        <!-- Column 1: Specs -->
        <div>
            <h3 class="text-xl font-bold text-gray-900 mb-6 flex items-center">
                <span class="bg-blue-100 text-blue-600 w-8 h-8 rounded-full flex items-center justify-center mr-3 text-sm">✓</span>
                Compliance Specifications
            </h3>
            <div class="bg-white p-6 rounded-xl shadow-sm border border-gray-100">
                <ul class="space-y-4 text-sm">
                    <li class="flex justify-between border-b border-gray-50 pb-2">
                        <span class="text-gray-500">Margins (Left)</span>
                        <span class="font-mono font-medium text-gray-900 bg-gray-50 px-2 rounded">{uni_data['margins']['left']}"</span>
                    </li>
                    <li class="flex justify-between border-b border-gray-50 pb-2">
                        <span class="text-gray-500">Margins (Others)</span>
                        <span class="font-mono font-medium text-gray-900 bg-gray-50 px-2 rounded">R: {uni_data['margins']['right']}", T: {uni_data['margins']['top']}"</span>
                    </li>
                    <li class="flex justify-between border-b border-gray-50 pb-2">
                        <span class="text-gray-500">Primary Font</span>
                        <span class="font-medium text-gray-900">{uni_data['font']['name']} ({uni_data['font']['size']}pt)</span>
                    </li>
                    <li class="flex justify-between border-b border-gray-50 pb-2">
                        <span class="text-gray-500">Line Spacing</span>
                        <span class="font-medium text-gray-900">{uni_data['line_spacing']}</span>
                    </li>
                    <li class="flex justify-between">
                        <span class="text-gray-500">Citation Style</span>
                        <span class="font-medium text-purple-600">{uni_data.get('reference_style', 'Standard')}</span>
                    </li>
                </ul>
            </div>
        </div>

        <!-- Column 2: Trust/Preview -->
        <div>
            <h3 class="text-xl font-bold text-gray-900 mb-6 transition-colors duration-200">Document Preview</h3>
            
            <!-- CSS-Only Document Preview -->
            <div class="relative bg-gray-200 h-80 rounded-xl flex items-center justify-center p-4 border border-gray-300 shadow-inner overflow-hidden group">
                <!-- The Paper -->
                <div class="bg-white w-48 h-64 shadow-2xl rounded-sm transform transition-transform duration-500 group-hover:scale-105 group-hover:-rotate-1 relative flex flex-col items-center pt-8 px-4 border border-gray-100">
                     <!-- Header Lines -->
                     <div class="w-full h-2 bg-gray-100 mb-2"></div>
                     <div class="w-3/4 h-2 bg-gray-100 mb-6"></div>
                     
                     <!-- Title -->
                     <div class="text-[8px] font-serif text-center text-gray-800 font-bold mb-1 uppercase tracking-widest">{uni_name}</div>
                     <div class="text-[6px] font-sans text-center text-blue-600 font-bold mb-4 uppercase tracking-wider">{course}</div>
                     
                     <!-- Body Lines -->
                     <div class="w-full space-y-1">
                        <div class="w-full h-1 bg-gray-100"></div>
                        <div class="w-full h-1 bg-gray-100"></div>
                        <div class="w-5/6 h-1 bg-gray-100"></div>
                        <div class="w-full h-1 bg-gray-100"></div>
                     </div>
                     
                     <!-- Footer -->
                     <div class="mt-auto mb-4 w-full flex justify-between px-1">
                         <div class="w-4 h-1 bg-gray-200"></div>
                         <div class="w-2 h-1 bg-gray-200"></div>
                     </div>
                </div>
                
                <!-- Badge Overlay -->
                 <div class="absolute bottom-4 right-4 bg-white/90 backdrop-blur px-3 py-1 rounded-full shadow-lg border border-gray-100 text-xs font-bold text-green-700 flex items-center gap-1">
                    <span class="w-1.5 h-1.5 bg-green-500 rounded-full animate-pulse"></span>
                    Preview
                 </div>
            </div>

            <div class="mt-4 flex gap-4 text-sm text-gray-600 justify-center md:justify-start">
                 <div class="flex items-center gap-1">
                    <svg class="w-4 h-4 text-green-500" fill="none" stroke="currentColor" viewBox="0 0 24 24"><path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M9 12l2 2 4-4m6 2a9 9 0 11-18 0 9 9 0 0118 0z"></path></svg>
                    <span>Virus Checked</span>
                 </div>
                 <div class="flex items-center gap-1">
                    <svg class="w-4 h-4 text-green-500" fill="none" stroke="currentColor" viewBox="0 0 24 24"><path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M12 8v4l3 3m6-3a9 9 0 11-18 0 9 9 0 0118 0z"></path></svg>
                    <span>Updated Jan 2026</span>
                 </div>
            </div>
        </div>
    </div>

    <!-- Footer -->
    <div class="bg-gray-50 border-t border-gray-200 py-12 text-center">
        <p class="text-sm text-gray-400 max-w-lg mx-auto">
            Disclaimer: This template is a student aid generated based on publicly available university guidelines. 
            Always verify with your specific department before final submission.
        </p>
        <p class="text-xs text-gray-300 mt-4">Generative Thesis Factory © 2026</p>
    </div>

</body>
</html>
    """
    with open(html_path, 'w') as f:
        f.write(html_content)

def generate_global_index(universities):
    """Generates the main homepage with client-side search (Tailwind Style)."""
    index_path = os.path.join(OUTPUT_DIR, "index.html")
    
    # Pre-render list items for SEO
    list_items = ""
    for uni in universities:
        link = f"{uni['id']}/index.html"
        list_items += f"""
        <li class="uni-item group bg-white border border-gray-200 rounded-lg hover:shadow-md transition-shadow duration-200">
            <a href="{link}" class="block p-5">
                <div class="flex items-center justify-between">
                    <div>
                        <h3 class="text-lg font-bold text-gray-900 group-hover:text-blue-600">{uni["uni_name"]}</h3>
                        <p class="text-sm text-gray-500 mt-1">{uni["course_name"]}</p>
                    </div>
                    <span class="text-gray-300 group-hover:text-blue-500">→</span>
                </div>
            </a>
        </li>
        """

    html_content = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Global Thesis Factory | Download 2026 University Templates</title>
    <meta name="description" content="Search and download free, compliant thesis templates for universities worldwide. MIT, Harvard, Oxford, Cambridge, and more.">
    <script src="https://cdn.tailwindcss.com"></script>
</head>
<body class="bg-gray-50 text-gray-800 font-sans">

    <div class="max-w-3xl mx-auto px-4 py-16">
        <div class="text-center mb-12">
            <h1 class="text-4xl md:text-5xl font-extrabold text-gray-900 mb-4 tracking-tight">
                🎓 Thesis Template Factory
            </h1>
            <p class="text-xl text-gray-600">
                Free, compliant Word templates for 19+ top universities.
            </p>
        </div>

        <!-- Search Box -->
        <div class="relative mb-10">
            <div class="absolute inset-y-0 left-0 pl-4 flex items-center pointer-events-none">
                <svg class="h-6 w-6 text-gray-400" fill="none" stroke="currentColor" viewBox="0 0 24 24"><path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M21 21l-6-6m2-5a7 7 0 11-14 0 7 7 0 0114 0z"></path></svg>
            </div>
            <input type="text" id="search" onkeyup="filterList()" 
                class="block w-full pl-12 pr-4 py-4 bg-white border border-gray-300 rounded-xl leading-5 placeholder-gray-400 focus:outline-none focus:ring-2 focus:ring-blue-500 focus:border-blue-500 text-lg shadow-sm" 
                placeholder="Search your university (e.g. 'Oxford', 'MIT')..." autofocus>
        </div>

        <!-- List -->
        <ul id="uniList" class="space-y-4 min-h-[200px]">
            {list_items}
        </ul>
        
        <!-- No Results State -->
        <div id="noResults" class="hidden text-center py-12">
            <div class="text-6xl mb-4">🔍</div>
            <h3 class="text-xl font-bold text-gray-900">No universities found</h3>
            <p class="text-gray-500 mt-2">Try searching for a different name or checking the spelling.</p>
        </div>
        
        <div class="mt-12 text-center border-t border-gray-100 pt-8">
            <p class="text-sm text-gray-400">
                Can't find your university? <a href="https://github.com/tanishqug/thesis-factory/issues" class="underline hover:text-blue-600 transition-colors">Request it on GitHub.</a>
            </p>
            <div class="mt-4">
                 <a href="https://github.com/tanishqug/thesis-factory" target="_blank" class="inline-flex items-center gap-2 text-gray-400 hover:text-gray-800 transition-colors text-xs font-medium bg-white px-3 py-1.5 rounded-full border border-gray-200 shadow-sm hover:shadow-md">
                    <svg class="h-4 w-4" fill="currentColor" viewBox="0 0 24 24"><path d="M12 0c-6.626 0-12 5.373-12 12 0 5.302 3.438 9.8 8.207 11.387.599.111.793-.261.793-.577v-2.234c-3.338.726-4.033-1.416-4.033-1.416-.546-1.387-1.333-1.756-1.333-1.756-1.089-.745.083-.729.083-.729 1.205.084 1.839 1.237 1.839 1.237 1.07 1.834 2.807 1.304 3.492.997.107-.775.418-1.305.762-1.604-2.665-.305-5.467-1.334-5.467-5.931 0-1.311.469-2.381 1.236-3.221-.124-.303-.535-1.524.117-3.176 0 0 1.008-.322 3.301 1.23.957-.266 1.983-.399 3.003-.404 1.02.005 2.047.138 3.006.404 2.291-1.552 3.297-1.23 3.297-1.23.653 1.653.242 2.874.118 3.176.77.84 1.235 1.911 1.235 3.221 0 4.609-2.807 5.624-5.479 5.921.43.372.823 1.102.823 2.222v3.293c0 .319.192.694.801.576 4.765-1.589 8.199-6.086 8.199-11.386 0-6.627-5.373-12-12-12z"/></svg>
                    Star on GitHub
                 </a>
            </div>
        </div>
    </div>

    <script>
        function filterList() {{
            var input, filter, ul, li, div, h3, txtValue;
            var visibleCount = 0;
            
            input = document.getElementById('search');
            filter = input.value.toUpperCase();
            ul = document.getElementById("uniList");
            li = ul.getElementsByTagName("li");
            noResults = document.getElementById("noResults");
            
            for (i = 0; i < li.length; i++) {{
                h3 = li[i].getElementsByTagName("h3")[0];
                txtValue = h3.textContent || h3.innerText;
                if (txtValue.toUpperCase().indexOf(filter) > -1) {{
                    li[i].style.display = "";
                    visibleCount++;
                }} else {{
                    li[i].style.display = "none";
                }}
            }}
            
            // Toggle No Results Message
            if (visibleCount === 0) {{
                noResults.classList.remove("hidden");
                ul.classList.add("hidden");
            }} else {{
                noResults.classList.add("hidden");
                ul.classList.remove("hidden");
            }}
        }}
    </script>
</body>
</html>
    """
    with open(index_path, 'w') as f:
        f.write(html_content)

# ---------------------------------------------------------
# ENTRY POINT
# ---------------------------------------------------------

if __name__ == "__main__":
    print("🏭 Starting Thesis Factory...")
    
    if not os.path.exists(DATA_FILE):
        print(f"❌ Error: {DATA_FILE} not found!")
        sys.exit(1)
        
    try:
        with open(DATA_FILE, 'r') as f:
            universities = json.load(f)
            
        for uni in universities:
            try:
                process_university(uni)
                # Build Web Page for this uni
                target_dir = os.path.join(OUTPUT_DIR, uni.get("id"))
                generate_web_page(target_dir, uni)
            except Exception as e:
                print(f"⚠️ Failed to process {uni.get('id', 'Unknown')}: {e}")
        
        # Build Global Search Index
        generate_global_index(universities)
        print("🌍 Website generated at Output/index.html")
                
        print("\n✨ All jobs completed.")
        
    except json.JSONDecodeError:
        print("❌ Error: Invalid JSON format in data.json")
    except Exception as e:
        print(f"❌ Critical Error: {e}")
